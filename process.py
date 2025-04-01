import os
import zipfile
import pdfplumber
import re
import time
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import subprocess
import yaml
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def timed_function(func):
    """Decorator to measure function execution time."""
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        elapsed_time = end_time - start_time
        if elapsed_time > 2:
            print(f"⏱ {func.__name__} took {elapsed_time:.4f} seconds")
        return result
    return wrapper

@timed_function
def load_yaml(yaml_path):
    """Load YAML configuration and return structured data."""
    with open(yaml_path, "r", encoding="utf-8") as file:
        yaml_data = yaml.safe_load(file)
    return yaml_data

@timed_function
def clean_text(text):
    """Clean text while preserving line breaks."""
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        cleaned_line = re.sub(r'[^a-zA-Z0-9\s\n*()\-,.:;?!\'"]', '', line)
        cleaned_lines.append(cleaned_line)
    return "\n".join(cleaned_lines)

@timed_function
def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF, using pdftotext first, then pdfplumber, then OCR if needed."""
    
    # Ensure the work_files directory exists
    output_dir = "work_files"
    os.makedirs(output_dir, exist_ok=True)

    # Construct the output file path
    output_file_path = os.path.join(output_dir, os.path.basename(pdf_path) + ".txt")

    # Try using pdftotext first
    result = subprocess.run(['pdftotext', pdf_path, '-'], capture_output=True, text=True)
    text = result.stdout.strip()

    if text:
        print(f"✅ Extracted text using pdftotext: {text[:100]}...")
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return text  # If pdftotext works, return immediately

    print("⚠️ pdftotext failed, trying pdfplumber...")

    # Fallback to pdfplumber
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    if text:
        print(f"✅ Extracted text using pdfplumber: {text[:100]}...")
        text = text.strip()
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return text  # If pdfplumber works, return immediately

    print("⚠️ pdfplumber failed, performing OCR...")

    # Final fallback: Use OCR (slow)
    text = ""
    images = convert_from_path(pdf_path)
    for img in images:
        ocr_text = pytesseract.image_to_string(img, lang='eng', config='--oem 3 --psm 6')
        cleaned_text = ocr_text.strip()
        text += cleaned_text + "\n"

    text = text.strip()
    print(f"✅ Extracted text using OCR (cleaned): {text[:100]}...")

    # Write the final extracted text to the file
    with open(output_file_path, "w", encoding="utf-8") as f:
        f.write(text)

    return text

@timed_function
def extract_text_from_docx(docx_path):
    """Extract text from Word document with error handling."""
    try:
        doc = Document(docx_path)
        return "\n".join(para.text for para in doc.paragraphs) or ""
    except Exception as e:
        print(f"⚠️ Error extracting text from {docx_path}: {str(e)}")
        return ""

@timed_function
def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add formatted paragraph to document."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

@timed_function
def process_sections(doc, sections, level=2, extracted_text=""):
    """Recursively process document sections."""
    for section in sections:
        if 'section' in section:
            add_formatted_paragraph(doc, section['section'], style=f'Heading {level}', bold=True)
            
            if section['search_pattern'] in extracted_text and section['extract_text']:
                match = extract_matching_text(
                    extracted_text,
                    section['search_pattern'],
                    section['extract_pattern'],
                    section['message_template']
                )
                if match:
                    add_formatted_paragraph(doc, match, italic=True)
                else:
                    add_formatted_paragraph(doc, "No matching content found.", style='Intense Quote')
            else:
                add_formatted_paragraph(doc, f"No {section['section']} information found.", style='Intense Quote')
            
            if 'sections' in section:
                process_sections(doc, section['sections'], level+1, extracted_text)

@timed_function
def extract_matching_text(text, search_pattern, extract_pattern, message_template):
    """Extract and format text using combined search and extract patterns."""
    # Combine the search pattern with the extract pattern
    full_pattern = search_pattern + extract_pattern
    matches = re.search(full_pattern, text, re.IGNORECASE | re.DOTALL)
    
    if matches:
        extracted = {f"extracted_text_{i+1}": matches.group(i+1) 
                    for i in range(matches.lastindex)}
        return message_template.format(**extracted)
    return None

@timed_function
def generate_report(doc, yaml_data, extracted_text):
    """Generate report document with validation."""
    # Ensure extracted_text is never None
    extracted_text = extracted_text or ""
    
    # Add title and scope
    doc.add_heading(yaml_data['general']['title'], level=0)
    scope = yaml_data['general']['scope'][0]
    doc.add_heading(scope['heading'], level=1)
    doc.add_paragraph(scope['body'])
    
    # Process each document section
    for doc_section in yaml_data['docs']:
        doc.add_heading(doc_section['heading'], level=1)
        
        # Check if identifier exists in text
        identifier = doc_section.get('identifier', '')
        if identifier and identifier in extracted_text:
            doc.add_paragraph(doc_section['message_if_identifier_found'])
        else:
            continue  # Skip processing this file if identifier not found
        
        # Process questions
        for question in doc_section.get('questions', []):
            # Handle address extraction specifically
            print(f"🔍 Processing question: {question}")
            if 'address' in question:
                print(f"🔍 Processing address with pattern: {question['search_pattern']}")
                add_formatted_paragraph(doc, question['address'], style='Heading 2')
                
                if question.get('search_pattern') and question.get('extract_text', False):
                    # Search for the address pattern
                    if question['search_pattern'] in extracted_text:
                        print(f"✅ Found address pattern in text")
                        address = extract_matching_text(
                            extracted_text, 
                            question['search_pattern'],  # Added this line
                            question['extract_pattern'], 
                            question['message_template']
                        )
                        if address:
                            print(f"✅ Extracted address: {address}")
                            add_formatted_paragraph(doc, address, italic=True)
                        else:
                            print("⚠️ Address pattern found but couldn't extract details")
                            add_formatted_paragraph(doc, "Address found but details couldn't be extracted", style='Intense Quote')
                    else:
                        print(f"⚠️ Address pattern not found in text")
                        add_formatted_paragraph(doc, "No address information found", style='Intense Quote')
            
            # Process other sections
            if 'sections' in question:
                print(f"🔍 Processing {len(question['sections'])} sections")
                process_sections(doc, question['sections'], extracted_text=extracted_text)

@timed_function
def process_zip(zip_path, output_docx, yaml_path):
    """Process ZIP file with improved error handling."""
    try:
        output_folder = "output_files/unzipped_files"
        os.makedirs(output_folder, exist_ok=True)
        
        yaml_data = load_yaml(yaml_path)
        doc = Document()
        
        # Add title and scope ONLY ONCE at the beginning
        doc.add_heading(yaml_data['general']['title'], level=0)
        scope = yaml_data['general']['scope'][0]
        doc.add_heading(scope['heading'], level=1)
        doc.add_paragraph(scope['body'])
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(output_folder)
        
        for file_name in os.listdir(output_folder):
            file_path = os.path.join(output_folder, file_name)
            
            try:
                if file_name.endswith(".pdf"):
                    extracted_text = extract_text_from_pdf(file_path)
                elif file_name.endswith(".docx"):
                    extracted_text = extract_text_from_docx(file_path)
                else:
                    continue
                
                # Ensure we have text to process
                if not extracted_text.strip():
                    print(f"⚠️ Empty text extracted from {file_name}")
                    continue
                    
                # Process document content without adding title/scope again
                process_document_content(doc, yaml_data, extracted_text)
                
            except Exception as e:
                print(f"⚠️ Error processing {file_name}: {str(e)}")
                continue
        
        os.makedirs(os.path.dirname(output_docx), exist_ok=True)
        doc.save(output_docx)
        
    except Exception as e:
        print(f"❌ Critical error processing ZIP: {str(e)}")
        raise

@timed_function
def process_document_content(doc, yaml_data, extracted_text):
    """Process document content without adding title/scope."""
    # Ensure extracted_text is never None
    extracted_text = extracted_text or ""
    
    # Process each document section
    for doc_section in yaml_data['docs']:
        
        # Check if identifier exists in text
        identifier = doc_section.get('identifier', '')
        if identifier and identifier in extracted_text:
            doc.add_heading(doc_section['heading'], level=1)
            doc.add_paragraph(doc_section['message_if_identifier_found'])
        else:
            continue  # Skip processing this file if identifier not found
        
        # Process questions
        for question in doc_section.get('questions', []):
            # Handle address extraction specifically
            print(f"🔍 Processing question: {question}")
            if 'address' in question:
                print(f"🔍 Processing address with pattern: {question['search_pattern']}")
                add_formatted_paragraph(doc, question['address'], style='Heading 2')
                
                if question.get('search_pattern') and question.get('extract_text', False):
                    # Search for the address pattern
                    if question['search_pattern'] in extracted_text:
                        print(f"✅ Found address pattern in text")
                        address = extract_matching_text(
                            extracted_text, 
                            question['extract_pattern'], 
                            question['message_template']
                        )
                        if address:
                            print(f"✅ Extracted address: {address}")
                            add_formatted_paragraph(doc, address, italic=True)
                        else:
                            print("⚠️ Address pattern found but couldn't extract details")
                            add_formatted_paragraph(doc, "Address found but details couldn't be extracted", style='Intense Quote')
                    else:
                        print(f"⚠️ Address pattern not found in text")
                        add_formatted_paragraph(doc, "No address information found", style='Intense Quote')
            
            # Process other sections
            if 'sections' in question:
                print(f"🔍 Processing {len(question['sections'])} sections")
                process_sections(doc, question['sections'], extracted_text=extracted_text)

# Main execution
if __name__ == "__main__":
    input_folder = "input_files"
    yaml_config = "config.yaml"
    output_file = "output_files/processed_doc.docx"
    
    # Find ZIP file
    zip_file_path = next(
        (os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".zip")),
        None
    )
    
    if zip_file_path:
        print(f"📂 Found ZIP file: {zip_file_path}")
        process_zip(zip_file_path, output_file, yaml_config)
    else:
        print("❌ No ZIP file found in 'input_files' folder.")
